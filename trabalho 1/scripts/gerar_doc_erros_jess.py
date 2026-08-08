from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm


ROOT = Path(__file__).resolve().parents[1]
ASSETS = Path(
    r"C:\Users\jesus\.cursor\projects\c-Users-jesus-OneDrive-Documentos-GitHub-TrabalhoPos\assets"
)

IMAGES = {
    "chat_ia": ASSETS
    / "c__Users_jesus_AppData_Roaming_Cursor_User_workspaceStorage_d5f94cd5b65fc5e7255f2f9475cf5e7f_images_WhatsApp_Image_2026-06-18_at_11.40.36-5355cc6f-9b02-4397-8043-9a502a0872fb.png",
    "estoque_duplicado": ASSETS
    / "c__Users_jesus_AppData_Roaming_Cursor_User_workspaceStorage_d5f94cd5b65fc5e7255f2f9475cf5e7f_images_WhatsApp_Image_2026-06-19_at_11.08.19-732c2428-c5c2-4a25-8310-588afa219159.png",
    "pagina_404": ASSETS
    / "c__Users_jesus_AppData_Roaming_Cursor_User_workspaceStorage_d5f94cd5b65fc5e7255f2f9475cf5e7f_images_WhatsApp_Image_2026-06-19_at_11.12.18-c5ca0b76-87ce-4514-afb8-530db4c6b2ad.png",
}


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Cm(0.6)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_image(doc: Document, path: Path, caption: str, width_cm: float = 14.5) -> None:
    if not path.exists():
        return
    doc.add_picture(str(path), width=Cm(width_cm))
    caption_p = doc.add_paragraph(caption)
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main() -> None:
    output = ROOT / "Erros_e_melhorias_reportados_Jess.docx"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Cm(0.42)

    add_title(doc, "Erros e melhorias reportados por Jess")
    doc.add_paragraph("Documento consolidado com problemas identificados durante o uso da aplicacao.")
    doc.add_paragraph("Periodo dos relatos: 18/06/2026 e 19/06/2026.")

    doc.add_heading("1. Conversa com a IA", level=1)
    add_bullet(
        doc,
        "Ao pedir uma recomendacao adicional depois de a IA ja ter sugerido carros, a conversa retornou a mensagem: "
        "'Desculpe, ocorreu um erro ao processar sua mensagem. Por favor, tente novamente.'",
    )
    add_bullet(
        doc,
        "Cenario informado: a IA ja havia recomendado opcoes e, em seguida, foi perguntado novamente qual carro deveria ser comprado.",
    )
    add_bullet(
        doc,
        "Comportamento esperado: a IA deveria continuar o contexto da conversa e indicar uma recomendacao final, sem falhar.",
    )
    add_image(doc, IMAGES["chat_ia"], "Evidencia 1 - falha ao continuar a conversa com a IA.", width_cm=7.5)

    doc.add_heading("2. Usabilidade do campo de mensagem", level=1)
    add_bullet(
        doc,
        "Enquanto a IA estiver respondendo, o campo de texto poderia continuar habilitado para o usuario ja escrever a proxima mensagem.",
    )
    add_bullet(
        doc,
        "O botao de enviar pode permanecer bloqueado durante o processamento, mas a digitacao nao deveria ser interrompida.",
    )
    add_bullet(
        doc,
        "A barra de texto deveria manter o foco apos cada envio, evitando a necessidade de clicar novamente para voltar a digitar.",
    )

    doc.add_heading("3. Criacao automatica de anuncios", level=1)
    add_bullet(
        doc,
        "Ao criar um anuncio, a janela/modal de criacao automatica nem sempre fecha imediatamente.",
    )
    add_bullet(
        doc,
        "Se o usuario clicar novamente enquanto a janela continua aberta, varios anuncios iguais podem ser criados em sequencia.",
    )
    add_bullet(
        doc,
        "Foi relatado que cerca de 5 anuncios foram criados ate a janela finalmente fechar.",
    )
    add_bullet(
        doc,
        "Comportamento esperado: apos a confirmacao da criacao, a janela deve fechar imediatamente e bloquear cliques repetidos.",
    )
    add_image(
        doc,
        IMAGES["estoque_duplicado"],
        "Evidencia 2 - anuncios semelhantes/duplicados aparecendo na listagem de estoque.",
        width_cm=16,
    )

    doc.add_heading("4. Validacao de anuncios duplicados", level=1)
    add_bullet(
        doc,
        "O sistema esta aceitando carros repetidos.",
    )
    add_bullet(
        doc,
        "Sugestao: se o anuncio for criado com exatamente os mesmos dados de outro ja existente, o sistema deve avisar que o anuncio ja existe e impedir a duplicacao.",
    )

    doc.add_heading("5. Abertura dos anuncios", level=1)
    add_bullet(
        doc,
        "Ao clicar nos anuncios dos carros, foi encontrada uma tela de erro 404.",
    )
    add_bullet(
        doc,
        "Comportamento esperado: o clique deve abrir a pagina de detalhes correta do anuncio.",
    )
    add_image(doc, IMAGES["pagina_404"], "Evidencia 3 - tela 404 ao abrir anuncio.", width_cm=16)

    doc.add_heading("6. Resumo executivo", level=1)
    add_bullet(doc, "Bug funcional na continuidade da conversa com a IA.")
    add_bullet(doc, "Melhorias de UX no campo de mensagem e foco de digitacao.")
    add_bullet(doc, "Problema de controle no fluxo de criacao automatica de anuncios.")
    add_bullet(doc, "Falta de validacao contra anuncios duplicados.")
    add_bullet(doc, "Erro de navegacao ao abrir anuncios (404).")

    doc.save(output)
    print(output)


if __name__ == "__main__":
    main()
